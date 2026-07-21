"""Office document loaders (DOCX, etc.)."""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from app.ingestion.loaders._common import tag_docs


def load_docx(path: Path) -> list[Document]:
    try:
        from langchain_community.document_loaders import Docx2txtLoader

        docs = Docx2txtLoader(str(path)).load()
    except Exception:
        docs = _load_docx_fallback(path)
    return tag_docs(docs, source_type="DOCX", content_type="TEXT")


def _load_docx_fallback(path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return [Document(page_content=text, metadata={"format": "docx"})] if text else []
